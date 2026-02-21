"""
Document processor for parsing and chunking files.
Supports PDF, DOCX, XLSX, CSV, TXT, and MD formats.
"""

from dataclasses import dataclass
from typing import List, BinaryIO
import io

from langchain_text_splitters import RecursiveCharacterTextSplitter


@dataclass
class DocumentChunk:
    """A chunk of text extracted from a document."""
    text: str
    chunk_index: int
    metadata: dict


class DocumentProcessor:
    """
    Handles document parsing and text chunking.
    
    Extracts text from various file formats and splits into
    overlapping chunks suitable for embedding and retrieval.
    """
    
    def __init__(
        self, 
        chunk_size: int = 500, 
        chunk_overlap: int = 50
    ):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Target size of each chunk in characters
            chunk_overlap: Number of overlapping characters between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def parse_file(
        self, 
        file_content: BinaryIO | bytes, 
        file_name: str,
        file_type: str
    ) -> str:
        """
        Extract text content from a file.
        
        Args:
            file_content: File content as bytes or file-like object
            file_name: Original file name
            file_type: File extension/type (pdf, docx, xlsx, csv, txt, md)
        
        Returns:
            Extracted text content
        """
        # Normalize file type
        file_type = file_type.lower().lstrip(".")
        
        # Ensure we have bytes
        if hasattr(file_content, "read"):
            content_bytes = file_content.read()
        else:
            content_bytes = file_content
        
        parsers = {
            "pdf": self._parse_pdf,
            "docx": self._parse_docx,
            "xlsx": self._parse_xlsx,
            "xls": self._parse_xlsx,
            "csv": self._parse_csv,
            "txt": self._parse_text,
            "md": self._parse_text,
            "markdown": self._parse_text,
        }
        
        parser = parsers.get(file_type)
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}. Supported: {list(parsers.keys())}")
        
        return parser(content_bytes)
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF file using pypdf."""
        from pypdf import PdfReader
        
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX file using python-docx."""
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n\n".join(text_parts)
    
    def _parse_xlsx(self, content: bytes) -> str:
        """Parse XLSX/XLS file using openpyxl."""
        from openpyxl import load_workbook
        
        wb = load_workbook(io.BytesIO(content), data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"## Sheet: {sheet_name}")
            
            for row in sheet.iter_rows():
                row_values = [str(cell.value) for cell in row if cell.value is not None]
                if row_values:
                    text_parts.append(" | ".join(row_values))
        
        return "\n".join(text_parts)
    
    def _parse_csv(self, content: bytes) -> str:
        """Parse CSV file."""
        import csv
        
        text = content.decode("utf-8", errors="ignore")
        reader = csv.reader(io.StringIO(text))
        rows = [" | ".join(row) for row in reader if any(row)]
        
        return "\n".join(rows)
    
    def _parse_text(self, content: bytes) -> str:
        """Parse plain text or markdown file."""
        return content.decode("utf-8", errors="ignore")
    
    def chunk_text(
        self, 
        text: str, 
        metadata: dict | None = None
    ) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Full text content to chunk
            metadata: Additional metadata to attach to each chunk
        
        Returns:
            List of DocumentChunk objects
        """
        if not text.strip():
            return []
        
        chunks = self.text_splitter.split_text(text)
        base_metadata = metadata or {}
        
        return [
            DocumentChunk(
                text=chunk,
                chunk_index=i,
                metadata={**base_metadata, "chunk_index": i, "total_chunks": len(chunks)}
            )
            for i, chunk in enumerate(chunks)
        ]
    
    def process_file(
        self,
        file_content: BinaryIO | bytes,
        file_name: str,
        file_type: str,
        additional_metadata: dict | None = None
    ) -> List[DocumentChunk]:
        """
        Parse a file and split into chunks in one step.
        
        Args:
            file_content: File content
            file_name: Original file name
            file_type: File extension
            additional_metadata: Extra metadata for chunks
        
        Returns:
            List of DocumentChunk objects ready for embedding
        """
        text = self.parse_file(file_content, file_name, file_type)
        
        metadata = {
            "file_name": file_name,
            "file_type": file_type,
            **(additional_metadata or {})
        }
        
        return self.chunk_text(text, metadata)
